import os
import pandas as pd
from docx import Document
from docx.shared import Inches

def generate_report():
    doc = Document()

    # Title
    doc.add_heading('Sales Analysis Report', 0)

    doc.add_heading('City and Branch Level Analysis', level=1)
    city_branch_sales_path = os.path.abspath(os.path.join('..', 'results', 'city_branch_sales.csv'))
    if os.path.exists(city_branch_sales_path):
        city_branch_sales = pd.read_csv(city_branch_sales_path)
        doc.add_paragraph('The following table summarizes the total sales at the city and branch level:')
        table = doc.add_table(city_branch_sales.shape[0] + 1, city_branch_sales.shape[1])
        table.style = 'Table Grid'
        
        for j, col in enumerate(city_branch_sales.columns):
            table.cell(0, j).text = col
        
        for i in range(city_branch_sales.shape[0]):
            for j in range(city_branch_sales.shape[1]):
                table.cell(i + 1, j).text = str(city_branch_sales.iat[i, j])
    else:
        doc.add_paragraph('City and Branch level sales data not found.')

    doc.add_page_break()

    doc.add_heading('Average Price of Items Sold at Each Branch', level=1)
    average_price_path = os.path.abspath(os.path.join('..', 'results', 'average_price.csv'))
    if os.path.exists(average_price_path):
        average_prices = pd.read_csv(average_price_path)
        doc.add_paragraph('The following table summarizes the average price of items sold at each branch:')
        table = doc.add_table(average_prices.shape[0] + 1, average_prices.shape[1])
        table.style = 'Table Grid'
        
        for j, col in enumerate(average_prices.columns):
            table.cell(0, j).text = col
        
        for i in range(average_prices.shape[0]):
            for j in range(average_prices.shape[1]):
                table.cell(i + 1, j).text = str(average_prices.iat[i, j])
    else:
        doc.add_paragraph('Average price data not found.')

    doc.add_page_break()

    doc.add_heading('Month-over-Month Analysis', level=1)
    month_over_month_path = os.path.abspath(os.path.join('..', 'results', 'month_over_month_analysis.csv'))
    if os.path.exists(month_over_month_path):
        monthly_analysis = pd.read_csv(month_over_month_path)
        doc.add_paragraph('The following table summarizes the month-over-month performance of sales and revenue:')
        table = doc.add_table(monthly_analysis.shape[0] + 1, monthly_analysis.shape[1])
        table.style = 'Table Grid'
        
        for j, col in enumerate(monthly_analysis.columns):
            table.cell(0, j).text = col
        
        for i in range(monthly_analysis.shape[0]):
            for j in range(monthly_analysis.shape[1]):
                table.cell(i + 1, j).text = str(monthly_analysis.iat[i, j])
    else:
        doc.add_paragraph('Month-over-month analysis data not found.')

    doc.add_heading('Visualizations', level=1)
    sales_trends_path = os.path.abspath(os.path.join('..', 'visualizations', 'sales_trends.png'))
    if os.path.exists(sales_trends_path):
        doc.add_paragraph('Sales Trends Over Time:')
        doc.add_picture(sales_trends_path, width=Inches(6))
    else:
        doc.add_paragraph('Sales trends visualization not found.')

    revenue_trends_path = os.path.abspath(os.path.join('..', 'visualizations', 'revenue_trends.png'))
    if os.path.exists(revenue_trends_path):
        doc.add_paragraph('Revenue Trends Over Time:')
        doc.add_picture(revenue_trends_path, width=Inches(6))
    else:
        doc.add_paragraph('Revenue trends visualization not found.')

    report_path = os.path.abspath(os.path.join('..', 'reports', 'sales_analysis_report.docx'))
    doc.save(report_path)

if __name__ == "__main__":
    generate_report()
